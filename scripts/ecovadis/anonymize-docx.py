#!/usr/bin/env python3
"""Anonymize EcoVadis DOCX outputs.

Replaces all mentions of ISLA variants with "Rated Company" and rewrites
"Ambito: <company>" into "Scope: Rated Company (GROUP)".
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document

ISLA_PATTERN = re.compile(r"\bISLA(?:\s*S\.?R\.?L\.?)?\b", re.IGNORECASE)
AMBITO_PATTERN = re.compile(r"\bAmbito\s*:\s*[^\n\r]+", re.IGNORECASE)


def normalize_text(value: str) -> str:
    text = AMBITO_PATTERN.sub("Scope: Rated Company (GROUP)", value)
    text = ISLA_PATTERN.sub("Rated Company", text)
    return text


def set_paragraph_text(paragraph, next_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(next_text)
        return

    paragraph.runs[0].text = next_text
    for run in paragraph.runs[1:]:
        run.text = ""


def rewrite_paragraph(paragraph) -> None:
    original = paragraph.text or ""
    updated = normalize_text(original)
    if updated != original:
        set_paragraph_text(paragraph, updated)


def rewrite_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                rewrite_paragraph(paragraph)
            for nested_table in cell.tables:
                rewrite_table(nested_table)


def rewrite_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        rewrite_paragraph(paragraph)

    for table in doc.tables:
        rewrite_table(table)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            rewrite_paragraph(paragraph)
        for table in section.header.tables:
            rewrite_table(table)

        for paragraph in section.footer.paragraphs:
            rewrite_paragraph(paragraph)
        for table in section.footer.tables:
            rewrite_table(table)


def assert_clean(doc: Document) -> None:
    fragments = [paragraph.text or "" for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fragments.extend(paragraph.text or "" for paragraph in cell.paragraphs)

    joined = "\n".join(fragments)
    if re.search(r"\bISLA\b", joined, re.IGNORECASE):
        raise RuntimeError("Anonymization failed: found residual 'ISLA' token")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Anonymize DOCX exports for EcoVadis delivery")
    parser.add_argument("input", type=Path, help="Input DOCX path")
    parser.add_argument("output", type=Path, nargs="?", help="Output DOCX path (default: overwrite input)")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = args.input
    output_path = args.output or args.input

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2

    doc = Document(str(input_path))
    rewrite_document(doc)
    assert_clean(doc)
    doc.save(str(output_path))

    print(f"Anonymized DOCX written to {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
